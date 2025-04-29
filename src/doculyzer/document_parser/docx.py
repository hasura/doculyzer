"""
DOCX document parser module for the document pointer system.

This module parses DOCX documents into structured elements.
"""

import hashlib
import json
import logging
import os
import uuid
from typing import Dict, Any, List, Optional

try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DocxDocument = None
    CT_Tbl = None
    CT_P = None
    Table = None
    Paragraph = None
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with 'pip install python-docx' to use DOCX parser")

from bs4 import BeautifulSoup

from .base import DocumentParser

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """Parser for DOCX documents."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize the DOCX parser."""
        super().__init__(config)

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")

        # Configuration options
        self.config = config or {}
        self.extract_comments = self.config.get("extract_comments", True)
        self.extract_headers_footers = self.config.get("extract_headers_footers", True)
        self.extract_styles = self.config.get("extract_styles", True)
        self.track_changes = self.config.get("track_changes", False)
        self.max_image_size = self.config.get("max_image_size", 1024 * 1024)  # 1MB default
        self.temp_dir = self.config.get("temp_dir", os.path.join(os.path.dirname(__file__), 'temp'))

    def parse(self, doc_content: Dict[str, Any]) -> Dict[str, Any]:
        """Parse a DOCX document into structured elements."""
        # Extract metadata from doc_content
        source_id = doc_content["id"]
        metadata = doc_content.get("metadata", {}).copy()  # Make a copy to avoid modifying original

        # Check if we have binary content or a path to a file
        binary_path = doc_content.get("binary_path")
        if not binary_path:
            # If we have binary content but no path, we need to save it to a temp file
            if not os.path.exists(self.temp_dir):
                os.makedirs(self.temp_dir, exist_ok=True)

            binary_content = doc_content.get("content", b"")
            if isinstance(binary_content, str):
                logger.warning("Expected binary content for DOCX but got string. Attempting to process anyway.")

            temp_file_path = os.path.join(self.temp_dir, f"temp_{uuid.uuid4().hex}.docx")
            with open(temp_file_path, 'wb') as f:
                if isinstance(binary_content, str):
                    f.write(binary_content.encode('utf-8'))
                else:
                    f.write(binary_content)

            binary_path = temp_file_path
            logger.debug(f"Saved binary content to temporary file: {binary_path}")

        # Generate document ID if not present
        doc_id = metadata.get("doc_id", self._generate_id("doc_"))

        # Load DOCX document
        try:
            doc = docx.Document(binary_path)
        except Exception as e:
            logger.error(f"Error loading DOCX document: {str(e)}")
            raise

        # Create document record with metadata
        document = {
            "doc_id": doc_id,
            "doc_type": "docx",
            "source": source_id,
            "metadata": self._extract_document_metadata(doc, metadata),
            "content_hash": doc_content.get("content_hash", "")
        }

        # Create root element
        elements = [self._create_root_element(doc_id, source_id)]
        root_id = elements[0]["element_id"]

        # Parse document elements
        elements.extend(self._parse_document_elements(doc, doc_id, root_id, source_id))

        # Extract links from the document
        links = self._extract_links(doc, elements)

        # Clean up temporary file if needed
        if binary_path != doc_content.get("binary_path") and os.path.exists(binary_path):
            try:
                os.remove(binary_path)
                logger.debug(f"Deleted temporary file: {binary_path}")
            except Exception as e:
                logger.warning(f"Failed to delete temporary file {binary_path}: {str(e)}")

        # Return the parsed document with extracted links
        return {
            "document": document,
            "elements": elements,
            "links": links,
            "relationships": []
        }

    def _extract_document_metadata(self, doc: DocxDocument, base_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: The DOCX document
            base_metadata: Base metadata from content source

        Returns:
            Dictionary of document metadata
        """
        # Combine base metadata with document core properties
        metadata = base_metadata.copy()

        try:
            # Get core properties
            core_props = doc.core_properties

            # Add core properties to metadata
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = core_props.created.timestamp() if hasattr(core_props.created,
                                                                                'timestamp') else str(
                    core_props.created)
            if core_props.modified:
                metadata["modified"] = core_props.modified.timestamp() if hasattr(core_props.modified,
                                                                                  'timestamp') else str(
                    core_props.modified)
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.category:
                metadata["category"] = core_props.category

            # Add document statistics
            metadata["page_count"] = self._estimate_page_count(doc)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["word_count"] = self._count_words(doc)

            # Add style information if enabled
            if self.extract_styles:
                metadata["styles"] = self._extract_styles_info(doc)

        except Exception as e:
            logger.warning(f"Error extracting document metadata: {str(e)}")

        return metadata

    def _parse_document_elements(self, doc: DocxDocument, doc_id: str, parent_id: str, source_id: str) -> List[
        Dict[str, Any]]:
        """
        Parse DOCX document into structured elements.

        Args:
            doc: The DOCX document
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            List of parsed elements
        """
        elements = []
        section_stack = [{"id": parent_id, "level": 0}]
        # current_parent = parent_id

        # Extract headers and footers if enabled
        if self.extract_headers_footers:
            header_elements = self._extract_headers_footers(doc, doc_id, parent_id, source_id)
            elements.extend(header_elements)

        # Process document body
        body_id = self._generate_id("body_")
        body_element = {
            "element_id": body_id,
            "doc_id": doc_id,
            "element_type": "body",
            "parent_id": parent_id,
            "content_preview": "Document body",
            "content_location": json.dumps({
                "source": source_id,
                "type": "body"
            }),
            "content_hash": "",
            "metadata": {}
        }
        elements.append(body_element)
        current_parent = body_id

        # Process all block-level elements in the document
        for i, block in enumerate(self._iter_block_items(doc)):
            if isinstance(block, Paragraph):
                # Process paragraph
                para_element = self._process_paragraph(block, i, doc_id, current_parent, source_id)

                # Skip empty paragraphs
                if not para_element:
                    continue

                # Check for headings
                style = block.style.name.lower() if block.style else ""
                if style.startswith('heading ') or style == 'title' or style == 'subtitle':
                    # This is a heading paragraph
                    level = 1  # Default level

                    if style == 'title':
                        level = 1
                    elif style == 'subtitle':
                        level = 2
                    elif style.startswith('heading '):
                        try:
                            level = int(style.split(' ')[1])
                        except (IndexError, ValueError):
                            pass

                    # Update element type and metadata
                    para_element["element_type"] = "header"
                    para_element["metadata"]["level"] = level

                    # Update section stack and current parent
                    while section_stack[-1]["level"] >= level:
                        section_stack.pop()

                    current_parent = section_stack[-1]["id"]
                    para_element["parent_id"] = current_parent

                    # Add to section stack
                    section_stack.append({"id": para_element["element_id"], "level": level})
                    current_parent = para_element["element_id"]

                elements.append(para_element)

            elif isinstance(block, Table):
                # Process table
                table_elements = self._process_table(block, i, doc_id, current_parent, source_id)
                elements.extend(table_elements)

        # Extract comments if enabled
        if self.extract_comments:
            try:
                comment_elements = self._extract_comments(doc, doc_id, body_id, source_id)
                elements.extend(comment_elements)
            except Exception as e:
                logger.warning(f"Error extracting comments: {str(e)}")

        return elements

    def _process_paragraph(self, paragraph: Paragraph, index: int, doc_id: str, parent_id: str, source_id: str) -> \
            Optional[Dict[str, Any]]:
        """
        Process a paragraph element.

        Args:
            paragraph: The paragraph
            index: Element index
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            Paragraph element dictionary or None if empty
        """
        # Get text content
        text = paragraph.text.strip()

        # Skip empty paragraphs
        if not text:
            return None

        # Generate element ID
        element_id = self._generate_id("para_")

        # Get paragraph style
        style_name = paragraph.style.name if paragraph.style else "Normal"

        # Get alignment
        alignment = "left"  # Default
        if paragraph.paragraph_format:
            if paragraph.paragraph_format.alignment:
                alignment_value = paragraph.paragraph_format.alignment
                if alignment_value == 1:
                    alignment = "center"
                elif alignment_value == 2:
                    alignment = "right"
                elif alignment_value == 3:
                    alignment = "justified"

        # Create paragraph element
        element = {
            "element_id": element_id,
            "doc_id": doc_id,
            "element_type": "paragraph",
            "parent_id": parent_id,
            "content_preview": text[:100] + ("..." if len(text) > 100 else ""),
            "content_location": json.dumps({
                "source": source_id,
                "type": "paragraph",
                "index": index
            }),
            "content_hash": self._generate_hash(text),
            "metadata": {
                "style": style_name,
                "alignment": alignment,
                "index": index
            }
        }

        # Check for list formatting
        if paragraph._p.pPr and paragraph._p.pPr.numPr:
            # This is a list item
            element["element_type"] = "list_item"

            # Try to determine list type and level
            list_level = 0
            if paragraph._p.pPr.numPr.ilvl:
                list_level = int(paragraph._p.pPr.numPr.ilvl.val)

            element["metadata"]["list_level"] = list_level

            # List type is harder to determine reliably without full numbering definitions
            # For now, we'll make a guess based on the first character
            if text.startswith(('•', '○', '■', '●', '◦', '◆')):
                element["metadata"]["list_type"] = "unordered"
            elif text.strip()[0].isdigit() and text.strip()[1:3] in ('. ', '.) '):
                element["metadata"]["list_type"] = "ordered"
            else:
                element["metadata"]["list_type"] = "unknown"

        return element

    def _process_table(self, table: Table, index: int, doc_id: str, parent_id: str, source_id: str) -> List[
        Dict[str, Any]]:
        """
        Process a table element.

        Args:
            table: The table
            index: Element index
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            List of table-related elements
        """
        elements = []

        # Generate table element ID
        table_id = self._generate_id("table_")

        # Create table element
        table_element = {
            "element_id": table_id,
            "doc_id": doc_id,
            "element_type": "table",
            "parent_id": parent_id,
            "content_preview": f"Table with {len(table.rows)} rows and {len(table.columns)} columns",
            "content_location": json.dumps({
                "source": source_id,
                "type": "table",
                "index": index
            }),
            "content_hash": "",
            "metadata": {
                "rows": len(table.rows),
                "columns": len(table.columns),
                "index": index
            }
        }
        elements.append(table_element)

        # Process rows
        for row_idx, row in enumerate(table.rows):
            # Generate row element ID
            row_id = self._generate_id("row_")

            # Create row element
            row_element = {
                "element_id": row_id,
                "doc_id": doc_id,
                "element_type": "table_row",
                "parent_id": table_id,
                "content_preview": f"Row {row_idx + 1}",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": "table_row",
                    "table_index": index,
                    "row": row_idx
                }),
                "content_hash": "",
                "metadata": {
                    "row": row_idx
                }
            }
            elements.append(row_element)

            # Process cells
            for col_idx, cell in enumerate(row.cells):
                # Generate cell element ID
                cell_id = self._generate_id("cell_")

                # Get cell content
                cell_text = " ".join(p.text for p in cell.paragraphs).strip()

                # Create cell element
                cell_element = {
                    "element_id": cell_id,
                    "doc_id": doc_id,
                    "element_type": "table_cell",
                    "parent_id": row_id,
                    "content_preview": cell_text[:100] + ("..." if len(cell_text) > 100 else ""),
                    "content_location": json.dumps({
                        "source": source_id,
                        "type": "table_cell",
                        "table_index": index,
                        "row": row_idx,
                        "col": col_idx
                    }),
                    "content_hash": self._generate_hash(cell_text),
                    "metadata": {
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text
                    }
                }

                # Check if this is a header cell (first row)
                if row_idx == 0:
                    cell_element["element_type"] = "table_header"

                elements.append(cell_element)

        return elements

    def _extract_headers_footers(self, doc: DocxDocument, doc_id: str, parent_id: str, source_id: str) -> List[
        Dict[str, Any]]:
        """
        Extract headers and footers from document.

        Args:
            doc: The DOCX document
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            List of header/footer elements
        """
        elements = []

        try:
            # Create headers container
            headers_id = self._generate_id("headers_")
            headers_element = {
                "element_id": headers_id,
                "doc_id": doc_id,
                "element_type": "headers",
                "parent_id": parent_id,
                "content_preview": "Document headers",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": "headers"
                }),
                "content_hash": "",
                "metadata": {}
            }
            elements.append(headers_element)

            # Create footers container
            footers_id = self._generate_id("footers_")
            footers_element = {
                "element_id": footers_id,
                "doc_id": doc_id,
                "element_type": "footers",
                "parent_id": parent_id,
                "content_preview": "Document footers",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": "footers"
                }),
                "content_hash": "",
                "metadata": {}
            }
            elements.append(footers_element)

            # Process sections
            for sect_idx, section in enumerate(doc.sections):
                # Process headers
                for header_type in ['first_page_header', 'header', 'even_page_header']:
                    header = getattr(section, header_type)
                    if header and header.is_linked_to_previous is False:
                        # Extract content
                        header_text = ""
                        for paragraph in header.paragraphs:
                            header_text += paragraph.text + "\n"

                        header_text = header_text.strip()
                        if header_text:
                            # Create header element
                            header_id = self._generate_id("header_")
                            header_element = {
                                "element_id": header_id,
                                "doc_id": doc_id,
                                "element_type": "header",
                                "parent_id": headers_id,
                                "content_preview": header_text[:100] + ("..." if len(header_text) > 100 else ""),
                                "content_location": json.dumps({
                                    "source": source_id,
                                    "type": "header",
                                    "section": sect_idx,
                                    "header_type": header_type
                                }),
                                "content_hash": self._generate_hash(header_text),
                                "metadata": {
                                    "section": sect_idx,
                                    "header_type": header_type.replace('_', ' '),
                                    "text": header_text
                                }
                            }
                            elements.append(header_element)

                # Process footers
                for footer_type in ['first_page_footer', 'footer', 'even_page_footer']:
                    footer = getattr(section, footer_type)
                    if footer and footer.is_linked_to_previous is False:
                        # Extract content
                        footer_text = ""
                        for paragraph in footer.paragraphs:
                            footer_text += paragraph.text + "\n"

                        footer_text = footer_text.strip()
                        if footer_text:
                            # Create footer element
                            footer_id = self._generate_id("footer_")
                            footer_element = {
                                "element_id": footer_id,
                                "doc_id": doc_id,
                                "element_type": "footer",
                                "parent_id": footers_id,
                                "content_preview": footer_text[:100] + ("..." if len(footer_text) > 100 else ""),
                                "content_location": json.dumps({
                                    "source": source_id,
                                    "type": "footer",
                                    "section": sect_idx,
                                    "footer_type": footer_type
                                }),
                                "content_hash": self._generate_hash(footer_text),
                                "metadata": {
                                    "section": sect_idx,
                                    "footer_type": footer_type.replace('_', ' '),
                                    "text": footer_text
                                }
                            }
                            elements.append(footer_element)
        except Exception as e:
            logger.warning(f"Error extracting headers/footers: {str(e)}")

        return elements

    def _extract_comments(self, doc: DocxDocument, doc_id: str, parent_id: str, source_id: str) -> List[Dict[str, Any]]:
        """
        Extract comments from document.

        Args:
            doc: The DOCX document
            doc_id: Document ID
            parent_id: Parent element ID
            source_id: Source identifier

        Returns:
            List of comment elements
        """
        elements = []

        try:
            # Create comments container
            comments_id = self._generate_id("comments_")
            comments_element = {
                "element_id": comments_id,
                "doc_id": doc_id,
                "element_type": "comments",
                "parent_id": parent_id,
                "content_preview": "Document comments",
                "content_location": json.dumps({
                    "source": source_id,
                    "type": "comments"
                }),
                "content_hash": "",
                "metadata": {}
            }
            elements.append(comments_element)

            # Extract comments
            # This is a bit tricky as python-docx doesn't have a direct API for comments
            # We need to access the XML directly

            # Get comments part if it exists
            if doc.part.package.parts:
                for rel_type, parts in doc.part.package.rels.items():
                    if 'comments' in rel_type.lower():
                        for rel_id, rel in parts.items():
                            if hasattr(rel, 'target_part') and rel.target_part:
                                # Process comments
                                comments_xml = rel.target_part.blob
                                if comments_xml:
                                    soup = BeautifulSoup(comments_xml, 'xml')
                                    for i, comment in enumerate(soup.find_all('comment')):
                                        # Extract comment data
                                        comment_id = comment.get('id', '')
                                        author = comment.get('author', 'Unknown')
                                        date = comment.get('date', '')
                                        text = comment.get_text().strip()

                                        # Create comment element
                                        comment_element_id = self._generate_id("comment_")
                                        comment_element = {
                                            "element_id": comment_element_id,
                                            "doc_id": doc_id,
                                            "element_type": "comment",
                                            "parent_id": comments_id,
                                            "content_preview": text[:100] + ("..." if len(text) > 100 else ""),
                                            "content_location": json.dumps({
                                                "source": source_id,
                                                "type": "comment",
                                                "comment_id": comment_id
                                            }),
                                            "content_hash": self._generate_hash(text),
                                            "metadata": {
                                                "comment_id": comment_id,
                                                "author": author,
                                                "date": date,
                                                "text": text,
                                                "index": i
                                            }
                                        }
                                        elements.append(comment_element)
        except Exception as e:
            logger.warning(f"Error extracting comments: {str(e)}")

        return elements

    def _extract_links(self, doc: DocxDocument, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Extract links from document.

        Args:
            doc: The DOCX document
            elements: Document elements

        Returns:
            List of extracted links
        """
        links = []

        try:
            # Extract hyperlinks
            # This is a bit tricky as python-docx doesn't have a direct API for hyperlinks
            # We need to access the XML directly

            # Iterate through relationships to find hyperlinks
            rels = doc.part.rels
            hyperlink_rels = {rel_id: rel.target_ref for rel_id, rel in rels.items()
                              if
                              rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'}

            # Iterate through all elements to find paragraphs
            for element in elements:
                element_id = element["element_id"]
                element_type = element["element_type"]

                # Only process text elements (paragraphs, headers, cells)
                if element_type not in ("paragraph", "header", "table_cell", "table_header", "list_item"):
                    continue

                # Get element content
                content_preview = element.get("content_preview", "")

                # Try to find hyperlinks in the XML for this element
                # This is a simplified approach that may not catch all hyperlinks
                for rel_id, target in hyperlink_rels.items():
                    # Check if this hyperlink's text appears in the element
                    # This is not perfect but a reasonable approximation
                    if target in content_preview:
                        # Create link
                        links.append({
                            "source_id": element_id,
                            "link_text": target,  # We don't know the exact text
                            "link_target": target,
                            "link_type": "hyperlink"
                        })
        except Exception as e:
            logger.warning(f"Error extracting links: {str(e)}")

        return links

    @staticmethod
    def _extract_styles_info(doc: DocxDocument) -> Dict[str, Any]:
        """
        Extract information about styles used in the document.

        Args:
            doc: The DOCX document

        Returns:
            Dictionary of style information
        """
        styles_info = {}

        try:
            # Get all styles in the document
            styles = doc.styles

            # Extract information about paragraph styles
            paragraph_styles = {}
            for style in styles:
                if style.type == 1:  # Paragraph style
                    paragraph_styles[style.name] = {
                        "style_id": style.style_id,
                        "based_on": style.base_style.name if style.base_style else None,
                        "builtin": not style.style_id.startswith('s')
                    }

            styles_info["paragraph_styles"] = paragraph_styles

            # Count usage of styles
            style_usage = {}
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name if paragraph.style else "Default"
                style_usage[style_name] = style_usage.get(style_name, 0) + 1

            styles_info["style_usage"] = style_usage

        except Exception as e:
            logger.warning(f"Error extracting style information: {str(e)}")

        return styles_info

    @staticmethod
    def _iter_block_items(doc: DocxDocument):
        """
        Iterate through all block items (paragraphs and tables) in document.

        Args:
            doc: The DOCX document

        Yields:
            Paragraph or Table objects
        """
        # Use XML to get all block items in order
        body = doc._body._body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc._body)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc._body)

    @staticmethod
    def _estimate_page_count(doc: DocxDocument) -> int:
        """
        Estimate page count for document.

        Args:
            doc: The DOCX document

        Returns:
            Estimated page count
        """
        # This is just a rough estimation as actual page count depends on formatting
        # Assuming 250 words per page on average
        word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
        page_count = max(1, word_count // 250)
        return page_count

    @staticmethod
    def _count_words(doc: DocxDocument) -> int:
        """
        Count words in document.

        Args:
            doc: The DOCX document

        Returns:
            Word count
        """
        return sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)

    @staticmethod
    def _generate_hash(content: str) -> str:
        """
        Generate a hash of content for change detection.

        Args:
            content: Text content

        Returns:
            MD5 hash of content
        """
        return hashlib.md5(content.encode('utf-8')).hexdigest()

    @staticmethod
    def _generate_id(prefix: str = "") -> str:
        """
        Generate a unique ID for a document or element.

        Args:
            prefix: Optional prefix for the ID

        Returns:
            Unique ID string
        """
        return f"{prefix}{uuid.uuid4()}"
